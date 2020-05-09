# -*- coding:UTF-8 -*-
import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn                   # 定义style
from docx.enum.section import WD_ORIENT       # 新建文件默认为竖向，变更页面方面必须同时修改页面尺寸
                                              # 定义页面方向 PORTRAIT  竖向
                                              #            LANDSCAPE 横向
from docx.enum.text import WD_ALIGN_PARAGRAPH # 定义对齐格式 CENTER, LEFT, RIGHT
                                              #            JUSTIFY    两端对齐
                                              #            DISTRIBUTE 分散对齐

class SaveToWord():
    def __init__(self, content, outputDir):
        self.content   = content     # 输出文档内容, list格式
        self.outputDir = outputDir   # 输出文件夹

    def AppFontStyle(self, run, runstyle):
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), runstyle['FontName'])
        
        run.font.color.rgb = runstyle['FontColor']
        run.font.size      = Pt(runstyle['FontSize'])
        run.bold           = runstyle['FontBold']
        run.italic         = runstyle['FontItalic']
        run.underline      = runstyle['FontUnderline']

    def SaveParagraph(self):
        # 设置doc文件相关参数
        titleStyle = {                 # 标题字体参数
            'FontName': '方正小标宋_GBK', # 名称
            'FontColor': RGBColor(0, 0, 0),     # 颜色
            'FontSize': 22,            # 字号， 二号=22， 小二=18， 三号=16， 小三=15, 四号=14， 小四=12
            'FontBold': False,         # 粗体
            'FontItalic': False,       # 斜体
            'FontUnderline': False     # 下划线
        }
        textStyle = {                  # 正文字体参数
            'FontName': '方正仿宋_GBK', # 字体
            'FontColor': RGBColor(0, 0, 0), # 颜色
            'FontSize': 16,            # 字号， 二号=22， 小二=18， 三号=16， 小三=15, 四号=14， 小四=12
            'FontBold': False,         # 粗体
            'FontItalic': False,       # 斜体
            'FontUnderline': False     # 下划线
        }

        # 打开输出文件
        if not self.content[0]:
            print('Content[0] is empty!')
            sys.exit()
        outputPath = self.outputDir + str(self.content[0]) + '.docx'
        if os.path.exists(outputPath):
            rsfile = Document(outputPath)
        else:
            rsfile = Document()
            
        # 设置页面格式
        section = rsfile.sections[0]
        section.orientation   = WD_ORIENT.PORTRAIT # 页面方向
        section.top_margin    = Cm(3.7) # 页边距
        section.bottom_margin = Cm(3.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.6)
            
        # 正文文本输出
        # 排版需要，先在标题前后各增加一个空行
        self.content.insert(0, '')
        self.content.insert(2, '')
        for ip in range(len(self.content)):
            p = rsfile.add_paragraph()
            p.paragraph_format.left_indent  = 0      # 左侧缩进
            p.paragraph_format.right_indent = 0      # 右侧缩进
            p.paragraph_format.space_before = 0      # 段前间距
            p.paragraph_format.space_after  = 0      # 段后间距
            p.paragraph_format.line_spacing = Pt(28) # 行间距
            
            run = p.add_run(self.content[ip])
            if ip <= 2:    # 标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.AppFontStyle(run, titleStyle)
            else:          # 正文
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(textStyle['FontSize']*2) # 首行缩进 
                self.AppFontStyle(run, textStyle)

        rsfile.save(outputPath)

#if __name__ == "__main__":
#    content = ['关于恳请协调在北沿江高铁如皋境内增设站点的情况汇报', '尊敬的李主席：', '首先，感谢您一直以来对如皋经济社会发展的关心和支持！今天来信，专程恳请您在百忙之中关心协调在北沿江高铁如皋境内增设站点事宜。']
#    outputDir = 'F:\\code-test\\python\\SaveToWord\\'
#    SaveToWord(content, outputDir).SaveParagraph()